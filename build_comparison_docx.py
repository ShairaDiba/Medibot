"""Build a Word document comparing Baseline vs LightRAG vs MediBot.

Includes: overall metrics tables, per-type breakdown, Hit@k curve,
retrieval-vs-answer gap analysis, and side-by-side sample outputs.
"""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent

lightrag = json.loads((ROOT / "eval_lightrag_results.json").read_text(encoding="utf-8"))
baseline = json.loads((ROOT / "eval_baseline_results.json").read_text(encoding="utf-8"))
medibot = json.loads((ROOT / "eval_results.json").read_text(encoding="utf-8"))
medibot_samples = json.loads((ROOT / "medibot_samples.json").read_text(encoding="utf-8"))

lr_by_id = {q["qid"]: q for q in lightrag["per_query"]}
bl_by_id = {q["qid"]: q for q in baseline["per_query"]}
mb_by_id = {q["qid"]: q for q in medibot["per_query"]}

doc = Document()

# ---- base styles ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def fmt(v, nd=4):
    return f"{v:.{nd}f}" if isinstance(v, (int, float)) else str(v)

# =====================================================================
# Title
# =====================================================================
title = doc.add_heading("MediBot vs LightRAG vs Baseline — Evaluation Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Corpus: Gale Encyclopedia of Medicine (5th Ed.) — 1,399/1,802 docs indexed for LightRAG (78%)", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
para("Queries: 60 (57 scored + 3 edge cases)  |  LLM: openai/gpt-oss-120b  |  Date: 2026-08-11", italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
# 1. Overall metrics
# =====================================================================
h1("1. Overall Metrics")

ls, bs, ms = lightrag["summary"], baseline["summary"], medibot["summary"]

h2("1.1 Retrieval Metrics")
make_table(
    ["Metric", "Baseline (no retrieval)", "LightRAG (mix)", "MediBot (KG BFS)", "Winner"],
    [
        ["Hit@1", "—", fmt(ls["hit1"]), fmt(ms["hit1"]), "MediBot (+91.6%)"],
        ["Hit@3", "—", fmt(ls["hit3"]), fmt(ms["hit3"]), "MediBot (+73.3%)"],
        ["Hit@5", "—", fmt(ls["hit5"]), fmt(ms["hit5"]), "MediBot (+71.0%)"],
        ["MRR",   "—", fmt(ls["mrr"]),  fmt(ms["mrr"]),  "MediBot (+82.7%)"],
    ],
)

h2("1.2 Response Quality (ROUGE / BLEU vs reference answers)")
make_table(
    ["Metric", "Baseline", "LightRAG", "MediBot", "Winner"],
    [
        ["ROUGE-1", fmt(bs["rouge1"]), fmt(ls["rouge1"]), fmt(ms["rouge1"]), "LightRAG (+27.1% vs base)"],
        ["ROUGE-2", fmt(bs["rouge2"]), fmt(ls["rouge2"]), fmt(ms["rouge2"]), "LightRAG (+54.5%)"],
        ["ROUGE-L", fmt(bs["rougeL"]), fmt(ls["rougeL"]), fmt(ms["rougeL"]), "LightRAG (+38.4%)"],
        ["BLEU-1",  fmt(bs["bleu1"]),  fmt(ls["bleu1"]),  fmt(ms["bleu1"]),  "LightRAG (+57.0%)"],
        ["BLEU-2",  fmt(bs["bleu2"]),  fmt(ls["bleu2"]),  fmt(ms["bleu2"]),  "LightRAG (+60.9%)"],
        ["BLEU-4",  fmt(bs["bleu4"]),  fmt(ls["bleu4"]),  fmt(ms["bleu4"]),  "MediBot (+87.9% vs base)"],
    ],
)

h2("1.3 Latency")
make_table(
    ["System", "Avg time per query"],
    [
        ["Baseline (no retrieval)", f"{bs.get('avg_elapsed_sec', 0):.1f}s"],
        ["MediBot (KG BFS)", f"{ms['avg_elapsed_sec']:.1f}s"],
        ["LightRAG (mix)", f"{ls['avg_elapsed_sec']:.1f}s"],
    ],
)

# =====================================================================
# 2. By query type
# =====================================================================
h1("2. Hit@1 by Query Type")
make_table(
    ["Query Type", "n", "LightRAG Hit@1", "MediBot Hit@1", "LightRAG MRR", "MediBot MRR"],
    [
        ["Definition", 20, "0.700", "1.000", "0.700", "1.000"],
        ["Symptom",    25, "0.200", "0.640", "0.313", "0.765"],
        ["Mixed",      10, "0.500", "1.000", "0.500", "1.000"],
    ],
)

# =====================================================================
# 3. Sample outputs
# =====================================================================
h1("3. Sample Outputs (Full Responses)")

SAMPLES = [
    ("DEF-001", "Definition", "Diabetes mellitus"),
    ("DEF-002", "Definition", "Tuberculosis"),
    ("SYM-001", "Symptom", "Abscess / Appendicitis / Aortic valve insufficiency"),
    ("SYM-002", "Symptom", "Allergic rhinitis"),
    ("MIX-001", "Mixed", "Epilepsy"),
]

for qid, qtype, expected in SAMPLES:
    lr = lr_by_id.get(qid, {})
    bl = bl_by_id.get(qid, {})
    mb = mb_by_id.get(qid, {})
    mb_sample = medibot_samples.get(qid, {})

    h2(f"3.{SAMPLES.index((qid, qtype, expected)) + 1}  {qid} ({qtype})")
    para(f"Query: {lr.get('query', bl.get('query', ''))}", bold=True)
    para(f"Expected disease(s): {expected}")

    # retrieval table
    lr_top5 = lr.get("ranked_top5", [])
    mb_top5 = mb.get("ranked_top5", []) or mb_sample.get("ranked", [])
    make_table(
        ["System", "Top-1 Retrieved", "Hit@1", "ROUGE-1"],
        [
            ["Baseline", "— (no retrieval)", "—", fmt(bl.get("rouge1", 0))],
            ["LightRAG", lr_top5[0] if lr_top5 else "—",
             "✅" if lr.get("hit1") else "❌", fmt(lr.get("rouge1", 0))],
            ["MediBot", (mb_top5[0] if mb_top5 else "—"),
             "✅" if mb.get("hit1") else "❌", fmt(mb.get("rouge1", 0))],
        ],
    )

    # full responses
    para("Baseline response:", bold=True)
    para(bl.get("llm_response", "(not recorded)"), size=9)
    para("LightRAG response:", bold=True)
    para(lr.get("llm_response", "(not recorded)"), size=9)
    para("MediBot response:", bold=True)
    para(mb_sample.get("llm_response", "(not recorded)"), size=9)

# =====================================================================
# 4. Takeaways
# =====================================================================
h1("4. Key Takeaways")
for t in [
    "Retrieval precision: MediBot ≫ LightRAG. Structured KG traversal (BFS from symptom seeds + alias recovery) beats LightRAG's generic entity extraction by ~92% on Hit@1. LightRAG often retrieves related-but-wrong entities (Diabetes → Hypoglycemia, Malaria → Antimalarial drugs).",
    "Response text quality: LightRAG > Baseline > MediBot on ROUGE/BLEU — LightRAG's long grounded context produces more reference-like prose, even when its top-1 entity is wrong.",
    "Latency: Baseline (~5s) < MediBot (20s) < LightRAG (41s). LightRAG pays a heavy cost for keyword extraction + dual-level graph retrieval per query.",
    "Symptom queries are the differentiator: MediBot 0.64 vs LightRAG 0.20 Hit@1 — the disease-symptom matrix + alias table does the heavy lifting that generic chunk embeddings can't replicate.",
    "Conclusion: For this medical lookup task, MediBot's structured KG approach is the clear winner on retrieval accuracy; LightRAG's value is fluent grounded prose, not precision. A hybrid (MediBot retrieval + LightRAG-style grounded generation) would likely dominate both axes.",
]:
    doc.add_paragraph(t, style="List Bullet")

# =====================================================================
# 5. Hit@k curve
# =====================================================================
def norm(s):
    return " ".join(str(s).lower().split())


def first_hit_rank(q):
    exp = {norm(e) for e in q.get("expected_diseases", [])}
    for i, t in enumerate(q.get("ranked_top5", []), 1):
        if norm(t) in exp:
            return i
    return None


lr_scored = [q for q in lightrag["per_query"] if q.get("type") != "edge"]
n_scored = len(lr_scored)

h1("5. LightRAG Hit@k Curve (where does the correct answer appear?)")
para("Cumulative share of 55 scored queries where the expected disease appears within the top-k retrieved entities:")
hitk_rows = []
prev = 0
for k in range(1, 6):
    hits = sum(1 for q in lr_scored if (first_hit_rank(q) or 99) <= k)
    hitk_rows.append([f"Hit@{k}", f"{hits}/{n_scored}", f"{hits / n_scored:.1%}", f"+{hits - prev}"])
    prev = hits
miss_n = n_scored - prev
hitk_rows.append(["Miss (not in top-5)", f"{miss_n}/{n_scored}", f"{miss_n / n_scored:.1%}", "—"])
make_table(["Metric", "Queries", "Rate", "Gain"], hitk_rows)
para("Key observation: the curve is nearly flat after k=1 — Hit@2 adds zero, and ranks 3-4 recover only 5 more queries. "
     "LightRAG either nails the answer at rank 1 or fails structurally; increasing k does not rescue it.", italic=True)

# =====================================================================
# 6. Retrieval vs answer correctness gap
# =====================================================================
h1("6. Retrieval Missed, But Did the Response Still Answer Correctly?")

REFUSAL_MARKERS = ("i'm sorry", "i am sorry", "don't have enough information",
                   "do not have enough information", "does not include any details",
                   "cannot give an answer", "cannot determine")


def response_mentions_expected(q):
    resp = norm(q.get("llm_response", "")).replace("\u2019", "'").replace("\u2018", "'")
    for e in q.get("expected_diseases", []):
        en = norm(e).replace("\u2019", "'").replace("\u2018", "'")
        if re.search(r"\b" + re.escape(en) + r"\b", resp):
            return e
        core = en.rstrip("s")
        if len(core) > 4 and re.search(r"\b" + re.escape(core) + r"\b", resp):
            return e
    return None


def is_refusal(q):
    resp = norm(q.get("llm_response", ""))
    return any(m in resp for m in REFUSAL_MARKERS)


lr_misses = [q for q in lr_scored if first_hit_rank(q) is None]
n_miss = len(lr_misses)
named = [q for q in lr_misses if response_mentions_expected(q)]
named_genuine = [q for q in named if not is_refusal(q)]
named_refusal = [q for q in named if is_refusal(q)]
fully_wrong = [q for q in lr_misses if not response_mentions_expected(q)]

para(f"Of the {n_miss} queries where LightRAG's top-5 retrieval missed the expected disease:")
make_table(
    ["Category", "Count", "Share of misses"],
    [
        ["Response still names the correct disease (genuine answer)", f"{len(named_genuine)}", f"{len(named_genuine) / n_miss:.0%}"],
        ["Response names the disease only in a refusal ('I don't have enough information about X')", f"{len(named_refusal)}", f"{len(named_refusal) / n_miss:.0%}"],
        ["Response also wrong", f"{len(fully_wrong)}", f"{len(fully_wrong) / n_miss:.0%}"],
    ],
)

h2("6.1 Effective end-to-end correctness")
eff_retrieval = len(lr_scored) - n_miss
eff_total = eff_retrieval + len(named_genuine)
make_table(
    ["Measure", "LightRAG", "MediBot"],
    [
        ["Retrieval Hit@1 (entity ranking)", f"{eff_retrieval}/{n_scored} = {eff_retrieval / n_scored:.0%}", f"{ms['hit1']:.0%}"],
        ["Effective answer correctness (retrieval + LLM compensation)", f"~{eff_total}/{n_scored} = ~{eff_total / n_scored:.0%}", f"~{ms['hit1']:.0%} (grounded)"],
    ],
)

h2("6.2 Why LightRAG answers correctly despite bad retrieval")
for t in [
    "LLM parametric knowledge fills the gap: for well-known diseases (diabetes, tuberculosis, malaria), the query itself names the topic and gpt-oss-120b already 'knows' the definition — the retrieved context is almost decorative.",
    "Near-miss retrieval still provides the right neighborhood: retrieving 'Antituberculosis drugs' yields chunks about tuberculosis, so the generated text stays on-topic.",
    "Refusal honesty: in a few cases (e.g., DEF-004 epilepsy, DEF-007 dengue) the model admitted the context was insufficient rather than hallucinating — a desirable safety behavior, but still a retrieval failure.",
]:
    doc.add_paragraph(t, style="List Bullet")

h2("6.3 Implication for system design")
para("LightRAG's pipeline is effectively 'weak retrieval + strong LLM fallback'. Its apparent end-to-end correctness (~80%) depends entirely on the LLM's pretrained medical knowledge. "
     "MediBot is 'strong retrieval + grounded generation' — its correctness comes from the curated knowledge graph, not the model's memory. "
     "If a smaller or cheaper LLM were substituted, LightRAG's effective accuracy would collapse while MediBot's would hold. "
     "For a medical assistant that must be trustworthy independent of model size, grounded retrieval (MediBot's approach) is the more robust architecture.", italic=False)

out = ROOT / "MediBot_vs_LightRAG_vs_Baseline_Evaluation.docx"
doc.save(out)
print(f"Saved -> {out}")
